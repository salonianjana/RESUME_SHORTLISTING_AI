import fitz  # PyMuPDF
import docx

def extract_text(file):
    text = ""

    # If the argument is a path string, handle accordingly
    if isinstance(file, str):
        if file.endswith('.pdf'):
            with fitz.open(file) as pdf:
                for page in pdf:
                    text += page.get_text()
        elif file.endswith('.docx'):
            doc = docx.Document(file)
            for para in doc.paragraphs:
                text += para.text + "\n"
        else:
            text = ""
    else:
        # If it’s a file object (for future use)
        if file.filename.endswith('.pdf'):
            with fitz.open(stream=file.read(), filetype="pdf") as pdf:
                for page in pdf:
                    text += page.get_text()
        elif file.filename.endswith('.docx'):
            doc = docx.Document(file)
            for para in doc.paragraphs:
                text += para.text + "\n"
        else:
            text = ""

    return text
